from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- CONFIG ---
# Using the local logo.png file instead of the URL
logo_path = "logo.png"
filename = "BlueRibbon_Nextdoor_Playbook.docx"
brand_blue = RGBColor(14, 76, 146)  # from your site
towns = ["Limerick","Stowe","West Pottsgrove","Sanatoga","Lower Pottsgrove",
          "North Coventry","South Limerick","Kenilworth","Parker Ford","Boyertown",
          "Douglassville","Gilbertsville","Royersford"]

# --- DOC SETUP ---
doc = Document()

# Check if logo exists locally, if so add it
if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Inches(2.5))
else:
    print(f"Warning: Local image '{logo_path}' not found. Skipping image insertion.")

doc.add_heading("Blue Ribbon Gutters – Nextdoor Playbook", level=0)

p = doc.add_paragraph("Be the name neighbors trust in " + ", ".join(towns) + ".")
p.runs[0].font.size = Pt(12)

def add_step(title, content):
    doc.add_heading(title, level=1)
    para = doc.add_paragraph(content)
    para.runs[0].font.size = Pt(11)
    para.runs[0].font.color.rgb = RGBColor(0,0,0)

# --- STEPS ---
add_step("⭐ STEP 1 — Get More Faves & Recommendations",
"""Goal: Hit 5+ Nextdoor recommendations for the shiny Neighborhood Favorite badge.

How:
• Text or email happy customers in your towns.
• Ask them to give Blue Ribbon a quick thumbs‑up on Nextdoor.
• Hit 5+ and your posts show to 30% more people!""")

add_step("📍 STEP 2 — Claim Your Zip Sponsorship",
"""Buy the zip codes you work in. You’ll always appear first when neighbors search “gutters.”
Average cost: $32–$150 a month. Way cheaper than a print ad.
Checklist:
☐ Log into Nextdoor Business page
☐ Search Neighborhood Sponsorships
☐ Claim your zips""")

add_step("🌦️ STEP 3 — Post When Weather Hits",
"""When rain or snow is coming, drop a friendly neighbor post:
"Hey Limerick neighbors! Heavy rain on Thursday. Clear your downspouts— we’re doing checks in Gilbertsville and Royersford this week." """)

add_step("📸 STEP 4 — Show You’re Local",
"""Use real photos, not stock images.
☑ Truck parked by a local landmark
☑ Before/after shots
☑ Tag the street or town name""")

add_step("🔔 STEP 5 — Jump on Threads Fast",
"""Turn on alerts for 'gutter', 'leak', 'roof', 'recommendation.'
Reply quick:
"Hi [Name]! We’re Blue Ribbon Gutters here in Parker Ford. Just finished a job nearby—happy to stop by today!" """)

# --- COMPARISON TABLE ---
doc.add_heading("Bojako vs. Blue Ribbon – Quick Comparison", level=1)
table = doc.add_table(rows=5, cols=3)
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = ["Feature", "Bojako (Old Way)", "Blue Ribbon (Smart Way)"]
rows = [
  ("Visibility", "Rely on old reviews", "Sponsors local ZIPs"),
  ("Content", "Generic posts", "Weather alerts + local photos"),
  ("Response", "Waits for tags", "Jumps in fast"),
  ("Conversion", "Links to homepage", "Links to local pages"),
]
for feat, old, new in rows:
    row_cells = table.add_row().cells
    row_cells[0].text, row_cells[1].text, row_cells[2].text = feat, old, new

# --- CHECKLIST ---
doc.add_heading("Weekly Checklist", level=1)
checks = [
"Get at least 1 new Nextdoor Fave",
"Make 1 weather post",
"Share 1 real job photo",
"Reply to 1 neighbor thread",
"Keep sponsorship active",
]
for item in checks:
    p = doc.add_paragraph("☐ " + item)

# --- FOOTER ---
doc.add_paragraph(
"People hire the neighbor they see and hear from the most — not the one with the biggest sign. "
"Show up, share real work, and stay friendly.\n\n"
"Blue Ribbon Gutters – Serving " + ", ".join(towns) + ".\nhttps://blueribbongutters.com"
)

# --- SAVE ---
doc.save(filename)
print(f"✅ Saved '{filename}' in current folder.")